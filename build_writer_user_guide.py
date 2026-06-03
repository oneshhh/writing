from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("writer_guide_output")
DOCX_PATH = OUT_DIR / "Real_Write_Writer_User_Guide.docx"

SCREENSHOTS = {
    "dashboard_sidebar": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 125718.png"),
    "create_blank": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 125733.png"),
    "create_filled": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 125817.png"),
    "editor": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131430.png"),
    "submit_controls": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131435.png"),
    "submitted": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131506.png"),
    "approved": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131634.png"),
    "earnings": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131651.png"),
    "rework_list": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131725.png"),
    "rework_submit": Path(r"C:\Users\Dell\Pictures\Screenshots\Screenshot 2026-06-03 131734.png"),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths, header_fill="E8EEF5"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0]
    for idx, text in enumerate(headers):
        set_cell_shading(hdr.cells[idx], header_fill)
        set_cell_text(hdr.cells[idx], text, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            set_cell_text(cells[idx], text)
    set_table_widths(table, widths)
    return table


def set_picture_alt(inline_shape, title, desc):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", desc)


def add_captioned_image(doc, path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    set_picture_alt(inline, caption, caption)
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullet(doc, text, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_note(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    r.font.size = Pt(10.5)
    for line in lines:
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_footer_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    add_footer_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(10.5)
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.25


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Real Write")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 54, 42)

    p = doc.add_paragraph()
    r = p.add_run("Writer User Guide")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    r = p.add_run("How to create, write, submit, track, revise, and monitor payments for writer articles.")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_note(
        doc,
        "Who this guide is for",
        [
            "Use this guide if you write articles in the Real Write writer portal.",
            "The screenshots use the writer account shown in the system, but the workflow is the same for each writer login.",
        ],
    )

    add_table(
        doc,
        ["Area", "Use it for"],
        [
            ("Dashboard", "Create new articles and track article status."),
            ("Editor", "Write the article, add SEO tags, save drafts, and submit final work."),
            ("Earnings", "View paid amount, pending amount, approved article count, and payment history."),
            ("Notifications", "Check system messages and review updates when available."),
            ("Profile", "Review account details."),
        ],
        [2500, 6860],
    )
    doc.add_page_break()


def build_doc():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Writer Portal Overview", level=1)
    doc.add_paragraph(
        "After logging in as a writer, the left sidebar gives access to Dashboard, Earnings, Notifications, and Profile. "
        "The active section is highlighted. The role label at the bottom confirms that the current login is a writer account, and Logout ends the session."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["dashboard_sidebar"],
        "Screenshot: Writer Dashboard with sidebar navigation, article status counters, and the My Articles table.",
    )

    doc.add_heading("What to Check First", level=2)
    add_bullet(doc, "Use the Approved, Rejected, and Rework counters to understand the current state of your articles.")
    add_bullet(doc, "Use My Articles to see each article title, project, status, AI score, plagiarism score, updated time, and action button.")
    add_bullet(doc, "Use New article when you need to start a new writing task.")
    add_bullet(doc, "Use View to open articles that are submitted or approved. Use Edit when an article is in rework.")

    doc.add_heading("2. Create a New Article", level=1)
    doc.add_paragraph(
        "From the Writer Dashboard, click New article. The New Article page asks for the project, type, title, and short description."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["create_blank"],
        "Screenshot: New Article page before entering a title and short description.",
    )
    doc.add_heading("Create Article Fields", level=2)
    add_table(
        doc,
        ["Field", "What to enter"],
        [
            ("Project", "Select the project assigned for this article, such as New writing Project."),
            ("Type", "Choose the article type, such as news, based on the assignment."),
            ("Title", "Enter the article headline or working title."),
            ("Short description", "Enter a 1-3 line summary of what the article will cover."),
        ],
        [2100, 7260],
    )
    doc.add_paragraph()
    add_numbered(doc, "Select the correct Project.")
    add_numbered(doc, "Select the correct Type.")
    add_numbered(doc, "Enter a clear Title.")
    add_numbered(doc, "Write a short description that summarizes the article.")
    add_numbered(doc, "Click Create. The system opens the editor after article creation.")
    add_captioned_image(
        doc,
        SCREENSHOTS["create_filled"],
        "Screenshot: New Article page with a project, type, title, and short description filled in.",
    )

    doc.add_heading("3. Write in the Editor", level=1)
    doc.add_paragraph(
        "The Editor page is where the article is completed. It includes the article title, SEO tags, short description, and long description editor."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["editor"],
        "Screenshot: Editor page with title, SEO tags, short description, and the long description writing area.",
    )
    doc.add_heading("Editor Fields", level=2)
    add_table(
        doc,
        ["Editor area", "How writers should use it"],
        [
            ("Title", "Review or update the article title before submission."),
            ("SEO tags", "Enter relevant keywords separated by commas, for example tag1, tag2."),
            ("Short description", "Keep this as a short summary of the article topic."),
            ("Long description", "Write the full article here using the toolbar for headings, lists, formatting, links, and images where needed."),
            ("Word counter", "Use the word count near the editor to confirm that the article is long enough before submission."),
        ],
        [2300, 7060],
    )
    add_note(
        doc,
        "Writing quality checklist before submission",
        [
            "Make sure the title, short description, and full article all match the assigned topic.",
            "Add SEO tags that match the article topic instead of generic tags.",
            "Use headings, short paragraphs, and lists to make the article easy to read.",
            "Check grammar, factual accuracy, duplicate text, and missing sections before submitting.",
            "Avoid submitting very short articles. The dashboard can show too_short under AI or Plagiarism when there is not enough text to evaluate.",
        ],
    )

    doc.add_heading("4. Save Draft or Submit", level=1)
    doc.add_paragraph(
        "At the bottom of the editor, the system shows the current article state, the Save draft button, and the Submit button."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["submit_controls"],
        "Screenshot: Bottom editor controls showing draft, Save draft, Submit, and word count.",
        width=4.7,
    )
    add_table(
        doc,
        ["Button or label", "Meaning"],
        [
            ("draft", "The article is currently not submitted and can still be edited."),
            ("Save draft", "Saves your work without sending it for review. Use this when the article is incomplete."),
            ("Submit", "Sends the article for review. Use this only after the full article is ready."),
        ],
        [2300, 7060],
    )

    doc.add_heading("5. Track Submitted and Approved Articles", level=1)
    doc.add_paragraph(
        "After submitting, return to the Dashboard. The article appears in My Articles with status submitted. The table also shows AI and plagiarism scores when available."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["submitted"],
        "Screenshot: Dashboard after submission, showing submitted status and AI/plagiarism indicators.",
    )
    doc.add_paragraph(
        "When the article is approved, the Approved counter increases and the article status changes to approved. Approved articles appear with a View button."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["approved"],
        "Screenshot: Dashboard after one article is approved.",
    )

    doc.add_heading("Status and Score Reference", level=2)
    add_table(
        doc,
        ["Label", "What it means", "What the writer should do"],
        [
            ("draft", "The article is saved but not submitted.", "Continue editing or submit when ready."),
            ("submitted", "The article has been sent for review.", "Wait for review. Use View to inspect the submitted article."),
            ("approved", "The article passed review.", "No further writing action is needed unless instructed."),
            ("rework", "The article needs changes.", "Open it with Edit, review remarks if shown, make changes, and submit again."),
            ("rejected", "The article was not accepted.", "Open View and check the remark column or review feedback."),
            ("AI score", "System AI-content indicator when the article is long enough to analyze.", "Keep writing original, human-edited content."),
            ("Plagiarism score", "Duplicate-content indicator when the article is long enough to analyze.", "Rewrite copied sections and cite/source facts appropriately."),
            ("too_short", "The article may not have enough words for a reliable check.", "Add complete article content before relying on the score."),
        ],
        [1700, 3830, 3830],
    )

    doc.add_heading("6. Handle Rework", level=1)
    doc.add_paragraph(
        "If an article is returned for rework, the Dashboard Rework counter increases. In My Articles, the article status becomes rework and the action button changes to Edit."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["rework_list"],
        "Screenshot: My Articles table with a rework article and an Edit action.",
    )
    doc.add_heading("Rework Steps", level=2)
    add_numbered(doc, "Open the Dashboard.")
    add_numbered(doc, "Find the article marked rework.")
    add_numbered(doc, "Click Edit.")
    add_numbered(doc, "Read any remark or feedback available for the article.")
    add_numbered(doc, "Update the title, SEO tags, short description, or article body as required.")
    add_numbered(doc, "Click Save draft if you are not done, or Submit when the revised article is ready.")
    add_captioned_image(
        doc,
        SCREENSHOTS["rework_submit"],
        "Screenshot: Rework editor controls showing rework status, Save draft, and Submit.",
        width=4.9,
    )

    doc.add_heading("7. Check Earnings and Payments", level=1)
    doc.add_paragraph(
        "Open Earnings from the sidebar to review payment totals and payment history. The page shows this month's paid amount, pending amount, and approved article count."
    )
    add_captioned_image(
        doc,
        SCREENSHOTS["earnings"],
        "Screenshot: Earnings page with this month's paid amount, pending amount, approved article count, and payment history.",
    )
    add_table(
        doc,
        ["Earnings item", "Meaning"],
        [
            ("Paid this month", "Amount already marked as paid during the month."),
            ("Pending this month", "Amount created but not yet paid."),
            ("Approved articles", "Number of approved articles counted for the month."),
            ("Payment History", "Article-level payment entries with amount, payment status, created date, and paid date."),
            ("pending", "The payment is not paid yet."),
            ("Paid at", "Shows the payment date after payment is completed. A dash means no paid date is recorded yet."),
        ],
        [2500, 6860],
    )

    doc.add_heading("8. Practical Rules for Writers", level=1)
    add_bullet(doc, "Create one article per assigned topic so tracking, review, and payment stay clean.")
    add_bullet(doc, "Save draft regularly while writing so incomplete work is not lost.")
    add_bullet(doc, "Submit only after the long description contains the complete article.")
    add_bullet(doc, "Watch the dashboard status instead of creating a duplicate article for the same assignment.")
    add_bullet(doc, "If an article is rework, edit the existing article and resubmit it from the editor.")
    add_bullet(doc, "Use Earnings to monitor pending and paid amounts, but contact the admin team for payment questions or missing entries.")

    doc.add_heading("Quick Workflow Summary", level=2)
    add_table(
        doc,
        ["Step", "Action", "Expected result"],
        [
            ("1", "Dashboard > New article", "New Article form opens."),
            ("2", "Fill project, type, title, and short description", "Article information is ready."),
            ("3", "Click Create", "Editor opens."),
            ("4", "Write article and SEO tags", "Article body is complete."),
            ("5", "Save draft or Submit", "Draft is saved or article enters review."),
            ("6", "Check Dashboard status", "Article shows submitted, approved, rejected, or rework."),
            ("7", "If rework, click Edit and resubmit", "Revised article returns to review."),
            ("8", "Open Earnings", "Payment status and history can be checked."),
        ],
        [900, 3700, 4760],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End of guide")
    r.italic = True
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    missing = [str(path) for path in SCREENSHOTS.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing screenshots: " + ", ".join(missing))
    build_doc()
    print(DOCX_PATH)
